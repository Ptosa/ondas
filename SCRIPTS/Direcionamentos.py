import os
import pandas as pd
from docx import Document
from docxcompose.composer import Composer
import tkinter as tk
from tkinter import messagebox

mes = None

def submit_data():
    global mes
    user_input = entry.get()
    try:
        mes = str(user_input)
        root.destroy()
    except ValueError:
        messagebox.showerror("Erro", "Por favor, insira um número válido")

root = tk.Tk()
root.title("Meu Programa")

label = tk.Label(root, text="Deseja pegar os direcionamentos de que mês (somente as 3 primeiras letras)")
label.pack(pady=10)

entry = tk.Entry(root)
entry.pack(pady=10)

button = tk.Button(root, text="Enviar", command=submit_data)
button.pack(pady=10)


root.mainloop()

# Função para salvar os dados de cada ministério em arquivos Excel separados
def salvar_ministerios_em_arquivos(df, diretorio_saida):
    # Verificar se o diretório de saída existe, se não, criar
    if not os.path.exists(diretorio_saida):
        os.makedirs(diretorio_saida)
    
    # Listar todas as colunas que contêm os ministérios
    colunas_ministerio = ['Ministério 1', 'Ministério 2', 'Ministério 3']
    
    for coluna in colunas_ministerio:
        # Remover espaços em branco antes e depois dos nomes dos ministérios
        df[coluna] = df[coluna].str.strip()
        # Listar todos os valores únicos encontrados na coluna de ministério atual
        valores_unicos_ministerio = df[coluna].dropna().unique()

        # Iterar sobre os ministérios possíveis
        for ministerio in valores_unicos_ministerio:
            # Filtrar o DataFrame pelo ministério atual
            ministerio_df = df[df[coluna] == ministerio]
            
            # Construir o caminho completo para salvar o arquivo
            caminho_arquivo = os.path.join(diretorio_saida, f'{ministerio}.xlsx')
            
            # Salvar o DataFrame em um arquivo Excel
            ministerio_df.to_excel(caminho_arquivo, index=False)


def criar_documento_direcionamento(df, ministerio, caminho):

    if pd.notnull(ministerio):
        documento = Document()  # Crie um novo documento

        for linha in df.index:
            #print(df.columns)
            # Definir ou atribuir valores às variáveis usadas como referências
            nome = df.loc[linha, 'Nome']
            telefone = df.loc[linha, 'Contato']
            apto = df.loc[linha, 'APTO P/ SERVIR']
            faltaum = 'X' if pd.isnull(df.loc[linha, 1]) else ''
            faltadois = 'X' if pd.isnull(df.loc[linha, 2]) else ''
            faltatres = 'X' if pd.isnull(df.loc[linha, 3]) else ''
            faltaquatro = 'X' if pd.isnull(df.loc[linha, 4]) else ''
            ficha = df.loc[linha, 'FICHA']
            quatrosim = 'X' if df.loc[linha, 4] == 'Sim' else ''
            quatronao = 'X' if df.loc[linha, 4] == 'Não' else ''
            nove = 'X' if str(df.loc[linha, 'H']).strip() == '9' or str(df.loc[linha, 'H']).strip() == '9h' else ''
            seis = 'X' if str(df.loc[linha, 'H']).strip() == '18' or str(df.loc[linha, 'H']).strip() == '18h' else ''
            minis=df.loc[linha,'Ministério 1']
            mes2=df.loc[linha,'Unnamed: 4']

            referencias = {
                "AAAA": nome,
                "BBBB": str(telefone),
                "CC": apto,
                "DD": faltaum,
                "EE": faltadois,
                "FF": faltatres,
                "GG": faltaquatro,
                "HH": ficha,
                "II": quatrosim,
                "JJ": quatronao,
                "KK": nove,
                "LL": seis,
                "XX": minis
            }

            # Percorrer todos os parágrafos do modelo de documento
            for paragrafo in Document(caminho).paragraphs:  # Substitua `seu_modelo_de_documento` pelo seu modelo real
                novo_paragrafo = paragrafo.text
                # Percorrer todas as chaves no dicionário de referências
                for codigo, valor in referencias.items():
                    # Substituir o código pelo valor no texto do parágrafo
                    
                    novo_paragrafo = novo_paragrafo.replace(codigo, str(valor))
                # Adicione o novo parágrafo ao documento
                if mes2 == mes:
                    documento.add_paragraph(novo_paragrafo)

        # Combine o caminho da pasta com o nome do arquivo
        caminho_completo = os.path.join(diretorio_saida, 'docx', f"{ministerio}_documento.docx")

        # Salve o documento com as substituições
        documento.save(caminho_completo)

# Caminho da pasta contendo os arquivos Excel
caminho_da_pasta = r'C:\Users\Pedro Lustosa\projeto_igreja\ministerios'

# Caminho do arquivo Excel
caminho_do_arquivo = r'C:\Users\Pedro Lustosa\projeto_igreja\CHAMADA+CADASTRO(3).xlsx'

# Carregar o arquivo Excel
df = pd.read_excel(caminho_do_arquivo)

# Excluir as duas primeiras linhas
#df = df.iloc[2:]

# Diretório de saída para salvar as planilhas Excel
diretorio_saida = r'C:\Users\Pedro Lustosa\projeto_igreja\ministerios'  

# Salvar os ministérios em arquivos separados
salvar_ministerios_em_arquivos(df, diretorio_saida)

caminho_arquivo_word = r'C:\Users\Pedro Lustosa\projeto_igreja\Direcionamentos.docx'

# Criar documentos de direcionamento para cada ministério
for ministerio in df['Ministério 1'].unique():
    if pd.notnull(ministerio):
     criar_documento_direcionamento(df[df['Ministério 1'] == ministerio], ministerio, caminho_arquivo_word)